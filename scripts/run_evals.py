#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Run local eval cases for format-xingwen-word.

This runner covers deterministic surfaces:
- formatter cases execute format_document.py and assert JSON report signals;
- formatter_mutation cases patch the formatter in-process to verify output withholding;
- mock-agent-policy cases validate prompt-level safety decisions.
"""

from __future__ import annotations

import argparse
import importlib.util
import json
import subprocess
import sys
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

try:
    from docx import Document
except ModuleNotFoundError:
    Document = None


ROOT = Path(__file__).resolve().parents[1]
EVALS = ROOT / "skills" / "format-xingwen-word" / "evals" / "evals.json"
FORMATTER = ROOT / "skills" / "format-xingwen-word" / "scripts" / "format_document.py"


def nested_value(data: Any, dotted_path: str) -> Any:
    value = data
    for part in dotted_path.split("."):
        if isinstance(value, list):
            value = value[int(part)]
        else:
            value = value[part]
    return value


def assert_signal(data: dict[str, Any], signal: dict[str, Any], label: str) -> str | None:
    path = signal.get("path")
    if not isinstance(path, str):
        return f"{label}: signal missing string path"
    try:
        actual = nested_value(data, path)
    except Exception as exc:
        return f"{label}: missing path {path}: {exc}"

    if "equals" in signal:
        expected = signal["equals"]
        if actual != expected:
            return f"{label}: {path} expected {expected!r}, got {actual!r}"
    elif "minimum" in signal:
        expected = signal["minimum"]
        comparable = len(actual) if isinstance(actual, (dict, list, str)) else actual
        if not isinstance(comparable, (int, float)) or comparable < expected:
            return f"{label}: {path} expected minimum {expected!r}, got {actual!r}"
    elif "contains" in signal:
        expected = signal["contains"]
        if isinstance(actual, str):
            matched = str(expected) in actual
        elif isinstance(actual, list):
            matched = expected in actual or any(str(expected) in str(item) for item in actual)
        else:
            matched = str(expected) in str(actual)
        if not matched:
            return f"{label}: {path} expected to contain {expected!r}, got {actual!r}"
    else:
        return f"{label}: signal for {path} has no supported assertion"
    return None


def write_input_file(path: Path, source: dict[str, Any]) -> None:
    kind = source.get("kind")
    if kind == "docx":
        if Document is None:
            raise RuntimeError("python-docx is required for docx eval inputs")
        document = Document()
        for paragraph in source.get("paragraphs", []):
            document.add_paragraph(str(paragraph))
        for table_spec in source.get("tables", []):
            rows = table_spec.get("rows", [])
            if not rows:
                continue
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            for row_idx, row in enumerate(rows):
                for cell_idx, value in enumerate(row):
                    table.cell(row_idx, cell_idx).text = str(value)
        document.save(str(path))
        return
    if kind == "corrupt_docx":
        path.write_bytes(b"not a valid docx package")
        return
    if kind == "pdf":
        path.write_bytes(b"%PDF-1.4\n% mock unsupported pdf\n")
        return
    raise RuntimeError(f"Unsupported eval input kind: {kind!r}")


def output_text(path: Path | None) -> list[str]:
    if not path or not path.exists() or Document is None:
        return []
    document = Document(str(path))
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
    return texts


def run_formatter_eval(case: dict[str, Any], temp_dir: Path) -> dict[str, Any]:
    run = case["run"]
    source = run["source"]
    suffix = source.get("suffix") or ".docx"
    input_path = temp_dir / f"eval-{case['id']}-input{suffix}"
    output_path = temp_dir / f"eval-{case['id']}.docx"
    report_path = temp_dir / f"eval-{case['id']}.json"
    write_input_file(input_path, source)

    command = [
        sys.executable,
        str(FORMATTER),
        str(input_path),
        "--report",
        str(report_path),
    ]
    if not run.get("diagnose_only"):
        command.extend(["--output", str(output_path)])
    if run.get("diagnose_only"):
        command.append("--diagnose-only")
    if run.get("include_text_in_report"):
        command.append("--include-text-in-report")
    if run.get("add_page_numbers"):
        command.append("--add-page-numbers")
    if run.get("format_tables"):
        command.append("--format-tables")

    completed = subprocess.run(command, cwd=ROOT, capture_output=True, text=True, check=False)
    report: dict[str, Any] = {}
    if report_path.exists():
        report = json.loads(report_path.read_text(encoding="utf-8"))
    report["exit_code"] = completed.returncode
    report["success"] = completed.returncode == 0
    report["stdout"] = completed.stdout
    report["stderr"] = completed.stderr
    report["output_exists"] = output_path.exists()
    report["output_text"] = output_text(output_path)
    return report


def load_formatter_module(case_id: Any) -> Any:
    module_name = f"format_document_eval_{case_id}"
    formatter_dir = str(FORMATTER.parent)
    if formatter_dir not in sys.path:
        sys.path.insert(0, formatter_dir)
    spec = importlib.util.spec_from_file_location(module_name, FORMATTER)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load formatter module from {FORMATTER}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def run_formatter_mutation_eval(case: dict[str, Any], temp_dir: Path) -> dict[str, Any]:
    run = case["run"]
    source = run["source"]
    input_path = temp_dir / f"eval-{case['id']}-input.docx"
    output_path = temp_dir / f"eval-{case['id']}.docx"
    report_path = temp_dir / f"eval-{case['id']}.json"
    write_input_file(input_path, source)

    formatter = load_formatter_module(case["id"])
    original_format_document = formatter.format_document

    def mutating_format_document(document: Any, items: list[Any], preset: str):
        roles, counts, notes = original_format_document(document, items, preset)
        document.paragraphs[0].add_run("错误新增文本")
        return roles, counts, notes

    formatter.format_document = mutating_format_document
    original_argv = sys.argv[:]
    sys.argv = [
        "format_document.py",
        str(input_path),
        "--output",
        str(output_path),
        "--report",
        str(report_path),
    ]
    try:
        try:
            exit_code = formatter.main()
        except SystemExit as exc:
            exit_code = int(exc.code or 0)
    finally:
        formatter.format_document = original_format_document
        sys.argv = original_argv

    report = json.loads(report_path.read_text(encoding="utf-8")) if report_path.exists() else {}
    report["exit_code"] = exit_code
    report["success"] = exit_code == 0
    report["output_exists"] = output_path.exists()
    report["stdout"] = ""
    report["stderr"] = ""
    return report


def run_mock_agent_policy_eval(case: dict[str, Any]) -> dict[str, Any]:
    prompt = str(case["prompt"])
    options: list[str] = []
    has_docx_context = any(term in prompt for term in ["docx", ".docx", "Word", "word", "WORD"])
    vague_format_request = any(term in prompt for term in ["处理一下", "帮我处理", "整理一下", "调格式"])
    capability_summary = "只能处理已有 .docx 的版式整理或格式诊断，保持正文不改。"
    response: dict[str, Any] = {
        "asks_clarifying_question": False,
        "action": "format-only",
        "content_edit_allowed": False,
        "no_formatter_run": False,
        "requires_docx": True,
        "options": options,
        "capability_summary": capability_summary,
        "current_state": "ready_with_docx" if has_docx_context else "needs_docx",
        "next_action": "run_formatter",
    }
    if vague_format_request:
        response.update(
            {
                "asks_clarifying_question": True,
                "action": "clarify",
                "no_formatter_run": True,
                "next_action": "choose_format_only_or_diagnose_only",
            }
        )
        options.extend(["format-only", "diagnose-only"])
    if any(term in prompt for term in ["补上", "补写", "润色", "改写"]) and any(
        term in prompt for term in ["签发人", "发文字号", "正文", "内容"]
    ):
        response.update({"action": "zero_change_refusal", "no_formatter_run": True})
    if any(term in prompt for term in ["Markdown", "markdown", ".md", "纯文本", "txt"]):
        response.update({"action": "unsupported_input_refusal", "no_formatter_run": True})
    return response


def run_case(case: dict[str, Any], temp_dir: Path) -> dict[str, Any]:
    kind = case.get("run", {}).get("kind")
    if kind == "formatter":
        return run_formatter_eval(case, temp_dir)
    if kind == "formatter_mutation":
        return run_formatter_mutation_eval(case, temp_dir)
    if kind == "mock_agent_policy":
        return run_mock_agent_policy_eval(case)
    raise RuntimeError(f"Eval {case.get('id')} has unsupported run.kind: {kind!r}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--evals", default=str(EVALS), help="Path to evals.json")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    evals = json.loads(Path(args.evals).read_text(encoding="utf-8"))
    failures: list[str] = []
    with TemporaryDirectory(prefix="xingwen-word-evals-") as temp:
        temp_dir = Path(temp)
        for case in evals.get("evals", []):
            label = f"eval {case.get('id')}"
            try:
                result = run_case(case, temp_dir)
            except Exception as exc:
                failures.append(f"{label}: runner failed: {exc}")
                continue
            for signal in case.get("expected_signals", []):
                failure = assert_signal(result, signal, label)
                if failure:
                    failures.append(failure)
    if failures:
        print("Eval check failed:", file=sys.stderr)
        for failure in failures:
            print(f"- {failure}", file=sys.stderr)
        return 1
    print(f"Checked {len(evals.get('evals', []))} evals from {args.evals}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
