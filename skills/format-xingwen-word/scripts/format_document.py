#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Format existing Word documents with the fixed 行文检查版式.

The script is intentionally conservative: it preserves paragraph text, applies
role-based formatting, and keeps reports free of full paragraph content by
default.
"""

from __future__ import annotations

import argparse
from collections import Counter
import hashlib
import json
import os
import re
import sys
from pathlib import Path
from typing import Any

from xingwen_word.presets import COMMON_REPORT_ROLES, LIMITED_FORMAT_AREAS, PRESETS
from xingwen_word.formatting import (
    ALIGNMENTS,
    apply_page_setup,
    apply_style,
    get_or_add_child,
    safe_xpath,
    safe_xpath_count,
    xml_child_val,
)
from xingwen_word.reports import (
    build_report,
    content_status,
    document_grid_reference,
    write_report,
)
from xingwen_word.roles import (
    BULLET_RE,
    DATE_RE,
    SEPARATOR_RE,
    ParagraphItem,
    RoleCandidate,
    decide_role,
    manual_numbering_kind,
    manual_numbering_level,
    normalize_text,
    paragraph_role_candidates,
)

try:
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ModuleNotFoundError as exc:
    if exc.name != "docx":
        raise
    Document = None
    WD_SECTION_START = None
    WD_STYLE_TYPE = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    WD_TABLE_ALIGNMENT = None
    WD_ALIGN_PARAGRAPH = None
    WD_LINE_SPACING = None
    WD_TAB_ALIGNMENT = None
    WD_TAB_LEADER = None
    OxmlElement = None
    parse_xml = None
    qn = None
    Mm = None
    Pt = None
    RGBColor = None
    DOCX_IMPORT_ERROR: ModuleNotFoundError | None = exc
else:
    DOCX_IMPORT_ERROR = None

IS_CLI_ENTRYPOINT = __name__ == "__main__"


def pt_value(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return float(value.pt)
    except AttributeError:
        try:
            return float(value)
        except Exception:
            return None


def mm_value(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.mm), 2)
    except AttributeError:
        try:
            return round(float(value), 2)
        except Exception:
            return None


def bool_or_none(value: Any) -> bool | None:
    if value is None:
        return None
    return bool(value)


def enum_name(value: Any) -> str | None:
    if value is None:
        return None
    name = getattr(value, "name", None)
    if name:
        return str(name)
    return str(value)


def require_docx() -> None:
    if DOCX_IMPORT_ERROR is None:
        return
    requirements_path = Path(__file__).with_name("requirements.txt")
    raise SystemExit(
        "Missing dependency: python-docx.\n"
        f'Install script dependencies with: python -m pip install -r "{requirements_path}"\n'
        "Or install directly with: python -m pip install python-docx"
    )


def text_hash(text: str) -> str:
    return hashlib.sha256(text.strip().encode("utf-8")).hexdigest()[:12]


def exact_text_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:12]


def iter_table_paragraphs(document: Any) -> list[Any]:
    paragraphs = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def document_text_units(document: Any) -> list[dict[str, Any]]:
    units: list[dict[str, Any]] = []
    for idx, paragraph in enumerate(document.paragraphs):
        units.append(
            {
                "identity": f"document.paragraphs[{idx}]",
                "scope": "body_paragraph",
                "text": paragraph.text,
            }
        )
    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for paragraph_idx, paragraph in enumerate(cell.paragraphs):
                    units.append(
                        {
                            "identity": (
                                f"tables[{table_idx}].rows[{row_idx}]"
                                f".cells[{cell_idx}].paragraphs[{paragraph_idx}]"
                            ),
                            "scope": "table_cell_paragraph",
                            "text": paragraph.text,
                        }
                    )
    for section_idx, section in enumerate(document.sections):
        parts = {
            ("header", "default"): section.header,
            ("header", "first_page"): section.first_page_header,
            ("header", "even_page"): section.even_page_header,
            ("footer", "default"): section.footer,
            ("footer", "first_page"): section.first_page_footer,
            ("footer", "even_page"): section.even_page_footer,
        }
        for (part_type, part_name), part in parts.items():
            for paragraph_idx, paragraph in enumerate(part.paragraphs):
                units.append(
                    {
                        "identity": (
                            f"sections[{section_idx}].{part_type}.{part_name}"
                            f".paragraphs[{paragraph_idx}]"
                        ),
                        "scope": "header_footer_paragraph",
                        "text": paragraph.text,
                    }
                )
            units.extend(
                text_box_text_units(
                    part._element,
                    f"sections[{section_idx}].{part_type}.{part_name}",
                )
            )
    units.extend(text_box_text_units(document._element.body, "document"))
    units.extend(related_text_part_units(document))
    return units


def related_text_part_scope(reltype: str) -> str | None:
    if reltype.endswith("/comments"):
        return "comments_paragraph"
    if reltype.endswith("/footnotes"):
        return "footnotes_paragraph"
    if reltype.endswith("/endnotes"):
        return "endnotes_paragraph"
    return None


def related_text_part_units(document: Any) -> list[dict[str, Any]]:
    units = []
    for rel in document.part.rels.values():
        scope = related_text_part_scope(str(rel.reltype))
        if scope is None:
            continue
        try:
            part = rel.target_part
        except Exception:
            continue
        element = getattr(part, "element", None)
        if element is None:
            element = getattr(part, "_element", None)
        if element is None and parse_xml is not None:
            try:
                element = parse_xml(part.blob)
            except Exception:
                element = None
        if element is None:
            continue
        part_name = str(getattr(part, "partname", scope))
        paragraphs = safe_xpath(element, ".//w:p")
        if not paragraphs and hasattr(element, "iter"):
            paragraphs = list(element.iter(qn("w:p")))
        for paragraph_idx, paragraph in enumerate(paragraphs):
            units.append(
                {
                    "identity": f"{part_name}.paragraphs[{paragraph_idx}]",
                    "scope": scope,
                    "text": element_text(paragraph),
                }
            )
        units.extend(text_box_text_units(element, part_name))
    return units


def element_text(element: Any) -> str:
    texts = []
    text_nodes = [*safe_xpath(element, ".//w:t"), *safe_xpath(element, ".//w:delText")]
    if not text_nodes and hasattr(element, "iter"):
        text_nodes = [
            *list(element.iter(qn("w:t"))),
            *list(element.iter(qn("w:delText"))),
        ]
    for node in text_nodes:
        if node.text:
            texts.append(node.text)
    return "".join(texts)


def text_box_text_units(element: Any, identity_prefix: str) -> list[dict[str, Any]]:
    units = []
    for box_idx, text_box in enumerate(safe_xpath(element, ".//w:txbxContent")):
        paragraphs = safe_xpath(text_box, "./w:p") or safe_xpath(text_box, ".//w:p")
        if not paragraphs and hasattr(text_box, "iter"):
            paragraphs = list(text_box.iter(qn("w:p")))
        for paragraph_idx, paragraph in enumerate(paragraphs):
            units.append(
                {
                    "identity": (
                        f"{identity_prefix}.text_boxes[{box_idx}]"
                        f".paragraphs[{paragraph_idx}]"
                    ),
                    "scope": "text_box_paragraph",
                    "text": element_text(paragraph),
                }
            )
    return units


def document_text_fingerprint(document: Any) -> dict[str, Any]:
    units = document_text_units(document)
    scope_counts = Counter(unit["scope"] for unit in units)
    unit_fingerprints = [
        {
            "identity": unit["identity"],
            "scope": unit["scope"],
            "hash": exact_text_hash(unit["text"]),
            "length": len(unit["text"]),
        }
        for unit in units
    ]
    combined = hashlib.sha256(
        json.dumps(unit_fingerprints, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()[:16]
    return {
        "scope": (
            "body paragraphs, table-cell paragraphs, header/footer paragraphs, "
            "text-box paragraphs, comments, footnotes, and endnotes after input parsing"
        ),
        "scope_counts": dict(scope_counts),
        "unit_count": len(unit_fingerprints),
        "non_empty_unit_count": len([unit for unit in units if unit["text"].strip()]),
        "combined_hash": combined,
        "units": unit_fingerprints,
    }


def compare_text_fingerprints(
    before: dict[str, Any],
    after: dict[str, Any],
    generated_layout_elements: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    before_units = before.get("units", [])
    after_units = after.get("units", [])
    changed_indexes = []
    for idx, (before_unit, after_unit) in enumerate(zip(before_units, after_units), start=1):
        if (
            before_unit.get("identity") != after_unit.get("identity")
            or before_unit.get("hash") != after_unit.get("hash")
            or before_unit.get("length") != after_unit.get("length")
        ):
            changed_indexes.append(idx)
    if len(before_units) != len(after_units):
        changed_indexes.extend(range(min(len(before_units), len(after_units)) + 1, max(len(before_units), len(after_units)) + 1))

    generated_layout_elements = generated_layout_elements or []
    return {
        "text_changed": before.get("combined_hash") != after.get("combined_hash"),
        "text_unit_count_changed": before.get("unit_count") != after.get("unit_count"),
        "paragraph_or_table_text_unit_count_changed": before.get("unit_count") != after.get("unit_count"),
        "paragraph_order_changed": [
            unit.get("identity") for unit in before_units
        ] != [
            unit.get("identity") for unit in after_units
        ],
        "generated_missing_elements": bool(generated_layout_elements),
        "generated_layout_elements": generated_layout_elements,
        "before_unit_count": before.get("unit_count"),
        "after_unit_count": after.get("unit_count"),
        "before_non_empty_unit_count": before.get("non_empty_unit_count"),
        "after_non_empty_unit_count": after.get("non_empty_unit_count"),
        "before_scope_counts": before.get("scope_counts"),
        "after_scope_counts": after.get("scope_counts"),
        "before_text_hash": before.get("combined_hash"),
        "after_text_hash": after.get("combined_hash"),
        "changed_unit_indexes": changed_indexes[:20],
        "scope": before.get("scope"),
    }


def document_items(document: Any) -> list[ParagraphItem]:
    return [
        ParagraphItem(text=p.text, paragraph=p)
        for p in document.paragraphs
        if p.text and p.text.strip()
    ]


def extract_run_font(run: Any) -> str | None:
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        east_asia = rpr.rFonts.get(qn("w:eastAsia"))
        if east_asia:
            return east_asia
    return run.font.name


def extract_run_style(run: Any) -> dict[str, Any]:
    style: dict[str, Any] = {}
    font = extract_run_font(run)
    if font:
        style["font"] = font
    if run.font.size is not None:
        style["size"] = pt_value(run.font.size)
    if run.font.bold is not None:
        style["bold"] = bool(run.font.bold)
    if run.font.italic is not None:
        style["italic"] = bool(run.font.italic)
    if run.font.underline is not None:
        style["underline"] = bool(run.font.underline)
    if run.font.strike is not None:
        style["strikethrough"] = bool(run.font.strike)
    if run.font.superscript:
        style["vertical_align"] = "superscript"
    elif run.font.subscript:
        style["vertical_align"] = "subscript"
    if run.font.all_caps is not None:
        style["all_caps"] = bool(run.font.all_caps)
    if run.font.small_caps is not None:
        style["small_caps"] = bool(run.font.small_caps)
    if run.font.highlight_color is not None:
        style["highlight"] = enum_name(run.font.highlight_color)
    if run.font.color is not None and run.font.color.rgb is not None:
        style["color"] = str(run.font.color.rgb)
    char_spacing = xml_child_val(run._element, "./w:rPr/w:spacing")
    if char_spacing is not None:
        style["character_spacing_twips"] = char_spacing
    emphasis = xml_child_val(run._element, "./w:rPr/w:em")
    if emphasis is not None:
        style["emphasis_mark"] = emphasis
    vanish = safe_xpath_count(run._element, "./w:rPr/w:vanish")
    if vanish:
        style["hidden"] = True
    return style


def extract_paragraph_style(paragraph: Any) -> dict[str, Any]:
    fmt = paragraph.paragraph_format
    style: dict[str, Any] = {}
    if paragraph.style is not None:
        style["style_name"] = paragraph.style.name
        style["style_id"] = paragraph.style.style_id
    if paragraph.alignment is not None:
        reverse = {value: key for key, value in ALIGNMENTS.items()}
        style["align"] = reverse.get(paragraph.alignment)
    if fmt.left_indent is not None:
        style["left_indent"] = pt_value(fmt.left_indent)
    if fmt.right_indent is not None:
        style["right_indent"] = pt_value(fmt.right_indent)
    if fmt.first_line_indent is not None:
        style["first_indent"] = pt_value(fmt.first_line_indent)
    if fmt.line_spacing is not None:
        style["line"] = pt_value(fmt.line_spacing)
    if fmt.space_before is not None:
        style["space_before"] = pt_value(fmt.space_before)
    if fmt.space_after is not None:
        style["space_after"] = pt_value(fmt.space_after)
    if fmt.keep_together is not None:
        style["keep_together"] = bool(fmt.keep_together)
    if fmt.keep_with_next is not None:
        style["keep_with_next"] = bool(fmt.keep_with_next)
    if fmt.page_break_before is not None:
        style["page_break_before"] = bool(fmt.page_break_before)
    if fmt.widow_control is not None:
        style["widow_control"] = bool(fmt.widow_control)
    outline_level = xml_child_val(paragraph._element, "./w:pPr/w:outlineLvl")
    if outline_level is not None:
        style["outline_level"] = outline_level
    num_ids = safe_xpath(paragraph._element, "./w:pPr/w:numPr/w:numId")
    if num_ids:
        style["numbering"] = {
            "num_id": num_ids[0].get(qn("w:val")),
            "level": xml_child_val(paragraph._element, "./w:pPr/w:numPr/w:ilvl"),
        }
    tab_stops = []
    try:
        for tab in fmt.tab_stops:
            tab_stops.append(
                {
                    "position_pt": pt_value(tab.position),
                    "alignment": enum_name(tab.alignment),
                    "leader": enum_name(tab.leader),
                }
            )
    except Exception:
        tab_stops = []
    if tab_stops:
        style["tab_stops"] = tab_stops

    for run in paragraph.runs:
        if run.text.strip():
            style.update(extract_run_style(run))
            break
    return {k: v for k, v in style.items() if v is not None}


def compact_style(style: dict[str, Any]) -> dict[str, Any]:
    keys = [
        "style_name",
        "style_id",
        "font",
        "size",
        "bold",
        "italic",
        "underline",
        "strikethrough",
        "color",
        "highlight",
        "vertical_align",
        "character_spacing_twips",
        "emphasis_mark",
        "hidden",
        "align",
        "first_indent",
        "left_indent",
        "right_indent",
        "line",
        "space_before",
        "space_after",
        "keep_together",
        "keep_with_next",
        "page_break_before",
        "widow_control",
        "outline_level",
        "numbering",
        "tab_stops",
    ]
    return {key: style.get(key) for key in keys if style.get(key) is not None}


def style_signature(style: dict[str, Any]) -> str:
    return json.dumps(compact_style(style), ensure_ascii=False, sort_keys=True)


def page_diagnostics(document: Any, preset: str) -> dict[str, Any]:
    sections = []
    expected = PRESETS[preset]["_page"]
    grid_reference = document_grid_reference(preset)
    for idx, section in enumerate(document.sections):
        sect_pr = section._sectPr
        columns = safe_xpath(sect_pr, "./w:cols")
        doc_grid = safe_xpath(sect_pr, "./w:docGrid")
        gutter = getattr(section, "gutter", None)
        width_mm = mm_value(section.page_width)
        height_mm = mm_value(section.page_height)
        left_mm = mm_value(section.left_margin)
        right_mm = mm_value(section.right_margin)
        top_mm = mm_value(section.top_margin)
        bottom_mm = mm_value(section.bottom_margin)
        gutter_mm = mm_value(gutter)
        text_area_width = None
        text_area_height = None
        if width_mm is not None and left_mm is not None and right_mm is not None:
            text_area_width = round(width_mm - left_mm - right_mm - (gutter_mm or 0), 2)
        if height_mm is not None and top_mm is not None and bottom_mm is not None:
            text_area_height = round(height_mm - top_mm - bottom_mm, 2)
        grid_info = {
            "type": doc_grid[0].get(qn("w:type")),
            "line_pitch_twips": doc_grid[0].get(qn("w:linePitch")),
            "char_space_twips": doc_grid[0].get(qn("w:charSpace")),
        } if doc_grid else None
        page = {
            "index": idx + 1,
            "width_mm": width_mm,
            "height_mm": height_mm,
            "top_margin_mm": top_mm,
            "bottom_margin_mm": bottom_mm,
            "left_margin_mm": left_mm,
            "right_margin_mm": right_mm,
            "gutter_mm": gutter_mm,
            "text_area_width_mm": text_area_width,
            "text_area_height_mm": text_area_height,
            "header_distance_mm": mm_value(section.header_distance),
            "footer_distance_mm": mm_value(section.footer_distance),
            "different_first_page_header_footer": bool(section.different_first_page_header_footer),
            "section_start_type": enum_name(section.start_type),
            "column_count": columns[0].get(qn("w:num")) if columns else None,
            "column_space_twips": columns[0].get(qn("w:space")) if columns else None,
            "document_grid": grid_info,
            "document_grid_reference": grid_reference,
            "expected_page_setup": {
                key: expected.get(key)
                for key in [
                    "width_mm",
                    "height_mm",
                    "top_mm",
                    "bottom_mm",
                    "left_mm",
                    "right_mm",
                    "text_area_width_mm",
                    "text_area_height_mm",
                    "grid_line_count",
                    "grid_char_count",
                    "grid_type",
                    "line_pitch_twips",
                    "char_space_twips",
                ]
                if expected.get(key) is not None
            },
            "orientation": "landscape"
            if section.page_width and section.page_height and section.page_width > section.page_height
            else "portrait",
        }
        differences = []
        for key, expected_key in [
            ("top_margin_mm", "top_mm"),
            ("bottom_margin_mm", "bottom_mm"),
            ("left_margin_mm", "left_mm"),
            ("right_margin_mm", "right_mm"),
        ]:
            value = page.get(key)
            exp = expected.get(expected_key)
            if value is not None and exp is not None and abs(value - exp) > 1:
                differences.append({"field": key, "actual": value, "expected": exp})
        for key in ["text_area_width_mm", "text_area_height_mm"]:
            value = page.get(key)
            exp = expected.get(key)
            if value is not None and exp is not None and abs(value - exp) > 1:
                differences.append({"field": key, "actual": value, "expected": exp})
        if grid_info:
            for actual_key, expected_key in [
                ("type", "grid_type"),
                ("line_pitch_twips", "line_pitch_twips"),
                ("char_space_twips", "char_space_twips"),
            ]:
                exp = expected.get(expected_key)
                actual = grid_info.get(actual_key)
                if exp is not None and actual is not None and str(actual) != str(exp):
                    differences.append({"field": "document_grid." + actual_key, "actual": actual, "expected": exp})
        elif expected.get("grid_type") or expected.get("line_pitch_twips"):
            differences.append({"field": "document_grid", "actual": None, "expected": "present"})
        if page["width_mm"] and page["height_mm"]:
            if abs(page["width_mm"] - 210) > 2 or abs(page["height_mm"] - 297) > 2:
                differences.append({"field": "paper", "actual": f"{page['width_mm']}x{page['height_mm']}mm", "expected": "A4 210x297mm"})
        page["differences_from_preset"] = differences
        sections.append(page)
    return {"section_count": len(sections), "sections": sections}


PAGE_CHANGE_FIELDS = [
    "width_mm",
    "height_mm",
    "top_margin_mm",
    "bottom_margin_mm",
    "left_margin_mm",
    "right_margin_mm",
    "gutter_mm",
    "text_area_width_mm",
    "text_area_height_mm",
    "header_distance_mm",
    "footer_distance_mm",
    "orientation",
    "section_start_type",
    "column_count",
    "column_space_twips",
    "document_grid.type",
    "document_grid.line_pitch_twips",
    "document_grid.char_space_twips",
]


def nested_value(data: dict[str, Any], dotted_key: str) -> Any:
    value: Any = data
    for part in dotted_key.split("."):
        if not isinstance(value, dict):
            return None
        value = value.get(part)
    return value


def compare_page_diagnostics(before: dict[str, Any], after: dict[str, Any]) -> dict[str, Any]:
    before_sections = before.get("sections", [])
    after_sections = after.get("sections", [])
    changes: list[dict[str, Any]] = []
    for idx in range(max(len(before_sections), len(after_sections))):
        before_section = before_sections[idx] if idx < len(before_sections) else {}
        after_section = after_sections[idx] if idx < len(after_sections) else {}
        section_changes = []
        for field in PAGE_CHANGE_FIELDS:
            before_value = nested_value(before_section, field)
            after_value = nested_value(after_section, field)
            if before_value != after_value:
                section_changes.append(
                    {
                        "field": field,
                        "before": before_value,
                        "after": after_value,
                    }
                )
        if section_changes:
            changes.append({"section_index": idx + 1, "changes": section_changes})
    return {
        "section_count_before": before.get("section_count"),
        "section_count_after": after.get("section_count"),
        "changed": bool(changes) or before.get("section_count") != after.get("section_count"),
        "changes": changes,
    }


def field_instructions(element: Any) -> list[str]:
    instructions = []
    for node in safe_xpath(element, ".//w:instrText"):
        if node.text:
            instructions.append(node.text)
    for node in safe_xpath(element, ".//w:fldSimple"):
        instr = node.get(qn("w:instr"))
        if instr:
            instructions.append(instr)
    return instructions


def field_type(instruction: str) -> str:
    match = re.search(r"[A-Za-z]+", instruction.strip())
    return match.group(0).upper() if match else "UNKNOWN"


def field_category(field_name: str) -> str:
    field_name = field_name.upper()
    if field_name == "TOC":
        return "table_of_contents"
    if field_name == "PAGE":
        return "page_number"
    if field_name == "NUMPAGES":
        return "total_pages"
    if field_name in {"DATE", "TIME", "CREATEDATE", "SAVEDATE", "PRINTDATE"}:
        return "date_time"
    if field_name in {"REF", "PAGEREF", "NOTEREF"}:
        return "cross_reference"
    if field_name in {"HYPERLINK", "INCLUDEPICTURE", "INCLUDETEXT"}:
        return "external_reference"
    if field_name in {"AUTHOR", "FILENAME", "FILESIZE", "TITLE", "SUBJECT", "COMMENTS"}:
        return "document_property"
    return "other"


def field_type_counts(instructions: list[str]) -> dict[str, int]:
    return dict(Counter(field_type(instruction) for instruction in instructions))


def field_category_counts(instructions: list[str]) -> dict[str, int]:
    return dict(Counter(field_category(field_type(instruction)) for instruction in instructions))


def field_instruction_details(
    instructions: list[str],
    *,
    scope: str,
    limit: int = 20,
) -> list[dict[str, Any]]:
    details = []
    for idx, instruction in enumerate(instructions[:limit], start=1):
        name = field_type(instruction)
        details.append(
            {
                "index": idx,
                "scope": scope,
                "field_type": name,
                "category": field_category(name),
                "instruction_hash": text_hash(instruction),
                "instruction_length": len(instruction.strip()),
            }
        )
    return details


FIELD_UPDATE_RISK_RULES = [
    {
        "category": "table_of_contents",
        "field_types": {"TOC"},
        "warning": "TOC field detected; update fields in Word/WPS after formatting so the table of contents reflects current headings and page numbers.",
    },
    {
        "category": "page_numbers",
        "field_types": {"PAGE", "NUMPAGES"},
        "warning": "PAGE/NUMPAGES fields detected; update fields in Word/WPS after formatting so page numbering reflects the final layout.",
    },
    {
        "category": "date_time",
        "field_types": {"DATE", "TIME", "CREATEDATE", "SAVEDATE", "PRINTDATE"},
        "warning": "Date/time fields detected; update fields in Word/WPS after formatting if the displayed date/time should refresh.",
    },
    {
        "category": "cross_reference",
        "field_types": {"REF", "PAGEREF", "NOTEREF"},
        "warning": "Cross-reference fields detected; update fields in Word/WPS after formatting so references and referenced page numbers are current.",
    },
]


def field_update_risks(field_diagnostics: dict[str, Any]) -> list[dict[str, Any]]:
    type_counts = field_diagnostics.get("all", {}).get("type_counts", {})
    risks = []
    for rule in FIELD_UPDATE_RISK_RULES:
        matching_counts = {
            field_type: type_counts.get(field_type)
            for field_type in sorted(rule["field_types"])
            if type_counts.get(field_type)
        }
        if not matching_counts:
            continue
        risks.append(
            {
                "category": rule["category"],
                "field_type_counts": matching_counts,
                "warning": rule["warning"],
            }
        )
    return risks


def field_count(instructions: list[str], field_name: str) -> int:
    pattern = re.compile(rf"\b{re.escape(field_name)}\b", re.IGNORECASE)
    return len([instruction for instruction in instructions if pattern.search(instruction)])


def paragraph_has_field(paragraph: Any, field_name: str) -> bool:
    return field_count(field_instructions(paragraph._element), field_name) > 0


def header_footer_part_diagnostics(part: Any) -> dict[str, Any]:
    paragraphs = [p for p in part.paragraphs if p.text.strip()]
    instructions = field_instructions(part._element)
    return {
        "has_text": bool(paragraphs),
        "paragraph_count": len(paragraphs),
        "field_instruction_count": len(instructions),
        "field_type_counts": field_type_counts(instructions),
        "field_category_counts": field_category_counts(instructions),
        "page_field_count": field_count(instructions, "PAGE"),
        "num_pages_field_count": field_count(instructions, "NUMPAGES"),
        "has_page_number_field": field_count(instructions, "PAGE") > 0,
    }


def header_footer_diagnostics(document: Any) -> dict[str, Any]:
    sections = []
    for idx, section in enumerate(document.sections):
        header_parts = {
            "default": section.header,
            "first_page": section.first_page_header,
            "even_page": section.even_page_header,
        }
        footer_parts = {
            "default": section.footer,
            "first_page": section.first_page_footer,
            "even_page": section.even_page_footer,
        }
        header_details = {}
        footer_details = {}
        for name, part in header_parts.items():
            header_details[name] = header_footer_part_diagnostics(part)
        for name, part in footer_parts.items():
            footer_details[name] = header_footer_part_diagnostics(part)
        header_text = " ".join(
            p.text.strip()
            for part in header_parts.values()
            for p in part.paragraphs
            if p.text.strip()
        )
        footer_text = " ".join(
            p.text.strip()
            for part in footer_parts.values()
            for p in part.paragraphs
            if p.text.strip()
        )
        sections.append(
            {
                "index": idx + 1,
                "has_header_text": bool(header_text),
                "has_footer_text": bool(footer_text),
                "has_page_number_field": any(
                    details.get("has_page_number_field")
                    for details in [*header_details.values(), *footer_details.values()]
                ),
                "page_field_count": sum(
                    int(details.get("page_field_count", 0))
                    for details in [*header_details.values(), *footer_details.values()]
                ),
                "different_first_page_header_footer": bool(section.different_first_page_header_footer),
                "header_parts": header_details,
                "footer_parts": footer_details,
            }
        )
    return {"sections": sections}


def add_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run._r.append(instr)

    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    paragraph.add_run("1")

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def page_number_style(preset: str) -> dict[str, Any]:
    base_font = PRESETS[preset].get("body", {}).get("font", "仿宋_GB2312")
    if preset in {"formal", "checklist"}:
        base_font = "宋体"
    return {
        "font": base_font,
        "size": 14,
        "align": "center",
        "first_indent": 0,
        "line": 18,
        "space_before": 0,
        "space_after": 0,
    }


def iter_header_footer_paragraphs(document: Any) -> list[tuple[int, str, str, Any]]:
    rows = []
    for section_idx, section in enumerate(document.sections, start=1):
        parts = {
            ("header", "default"): section.header,
            ("header", "first_page"): section.first_page_header,
            ("header", "even_page"): section.even_page_header,
            ("footer", "default"): section.footer,
            ("footer", "first_page"): section.first_page_footer,
            ("footer", "even_page"): section.even_page_footer,
        }
        for (part_type, part_name), part in parts.items():
            for paragraph in part.paragraphs:
                rows.append((section_idx, part_type, part_name, paragraph))
    return rows


def format_existing_page_numbers(document: Any, preset: str) -> list[dict[str, Any]]:
    actions = []
    style = page_number_style(preset)
    for section_idx, part_type, part_name, paragraph in iter_header_footer_paragraphs(document):
        if paragraph_has_field(paragraph, "PAGE"):
            apply_style(paragraph, style)
            actions.append(
                {
                    "action": "formatted_existing_page_number",
                    "section": section_idx,
                    "part_type": part_type,
                    "part": part_name,
                }
            )
    return actions


def footer_parts_for_page_number(section: Any) -> list[tuple[str, Any]]:
    parts = [("default", section.footer)]
    if section.different_first_page_header_footer:
        parts.append(("first_page", section.first_page_footer))
    return parts


def part_has_page_number(part: Any) -> bool:
    return field_count(field_instructions(part._element), "PAGE") > 0


def add_page_numbers(document: Any, preset: str) -> list[dict[str, Any]]:
    actions = []
    style = page_number_style(preset)
    for section_idx, section in enumerate(document.sections, start=1):
        for part_name, footer in footer_parts_for_page_number(section):
            if part_has_page_number(footer):
                actions.append(
                    {
                        "action": "preserved_existing_page_number",
                        "section": section_idx,
                        "part_type": "footer",
                        "part": part_name,
                    }
                )
                continue
            paragraph = footer.paragraphs[0] if footer.paragraphs and not footer.paragraphs[0].text.strip() else footer.add_paragraph()
            paragraph.add_run("— ")
            add_field(paragraph, "PAGE")
            paragraph.add_run(" —")
            apply_style(paragraph, style)
            actions.append(
                {
                    "action": "added_page_number",
                    "section": section_idx,
                    "part_type": "footer",
                    "part": part_name,
                    "format": "— PAGE —",
                }
            )
    return actions


def set_table_width_pct(table: Any, pct: int = 5000) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = get_or_add_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct))


def set_table_borders(table: Any, color: str = "000000", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = get_or_add_child(tbl_pr, "w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = "w:" + edge
        border = get_or_add_child(borders, tag)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_cell_margins(cell: Any, margin_twips: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = get_or_add_child(tc_pr, "w:tcMar")
    for side in ["top", "left", "bottom", "right"]:
        node = get_or_add_child(tc_mar, "w:" + side)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = get_or_add_child(tr_pr, "w:tblHeader")
    header.set(qn("w:val"), "true")


def table_has_merged_cells(table: Any) -> bool:
    return bool(
        safe_xpath_count(table._tbl, ".//w:gridSpan")
        or safe_xpath_count(table._tbl, ".//w:vMerge")
    )


def table_structure_options(preset: str) -> dict[str, Any]:
    return {
        "width_pct": 5000,
        "border_color": "000000",
        "border_size": "4",
        "cell_margin_twips": 80,
        "repeat_first_row": True,
        "vertical_alignment": "center",
        "alignment": "center",
        "autofit": False,
    }


def format_table_structures(document: Any, preset: str) -> list[dict[str, Any]]:
    actions = []
    options = table_structure_options(preset)
    for idx, table in enumerate(document.tables, start=1):
        action: dict[str, Any] = {
            "action": "formatted_table_structure",
            "table_index": idx,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "merged_cells_detected": table_has_merged_cells(table),
            "applied": [],
        }
        if WD_TABLE_ALIGNMENT is not None:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            action["applied"].append("alignment_center")
        table.autofit = bool(options["autofit"])
        action["applied"].append("autofit_false")
        set_table_width_pct(table, int(options["width_pct"]))
        action["applied"].append("width_100_percent")
        set_table_borders(table, str(options["border_color"]), str(options["border_size"]))
        action["applied"].append("single_black_borders")
        if options["repeat_first_row"] and table.rows:
            set_repeat_header(table.rows[0])
            action["applied"].append("repeat_first_row")
        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell, int(options["cell_margin_twips"]))
                if WD_CELL_VERTICAL_ALIGNMENT is not None:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        action["applied"].append("cell_margins")
        if WD_CELL_VERTICAL_ALIGNMENT is not None:
            action["applied"].append("vertical_center")
        actions.append(action)
    return actions


def table_diagnostics(document: Any) -> dict[str, Any]:
    tables = []
    for idx, table in enumerate(document.tables):
        sample_styles = []
        row_heights = []
        column_widths = []
        cell_margins = []
        header_row_count = 0
        for row in table.rows[:2]:
            if row.height is not None:
                row_heights.append(mm_value(row.height))
            if safe_xpath_count(row._tr, "./w:trPr/w:tblHeader"):
                header_row_count += 1
            for cell in row.cells[:3]:
                if cell.width is not None:
                    column_widths.append(mm_value(cell.width))
                tc_mar = safe_xpath(cell._tc, "./w:tcPr/w:tcMar")
                if tc_mar:
                    margins = {}
                    for side in ["top", "bottom", "left", "right"]:
                        side_node = tc_mar[0].find(qn("w:" + side))
                        if side_node is not None:
                            margins[side] = side_node.get(qn("w:w"))
                    if margins:
                        cell_margins.append(margins)
                for p in cell.paragraphs:
                    if p.text.strip():
                        sample_styles.append(compact_style(extract_paragraph_style(p)))
                        break
        tbl_pr = table._tbl.tblPr
        tbl_width = safe_xpath(tbl_pr, "./w:tblW")
        tbl_borders = safe_xpath_count(tbl_pr, "./w:tblBorders/*")
        tables.append(
            {
                "index": idx + 1,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "alignment": enum_name(table.alignment),
                "autofit": bool(table.autofit),
                "merged_cell_indicators": {
                    "grid_span_count": safe_xpath_count(table._tbl, ".//w:gridSpan"),
                    "vertical_merge_count": safe_xpath_count(table._tbl, ".//w:vMerge"),
                },
                "width_twips": tbl_width[0].get(qn("w:w")) if tbl_width else None,
                "width_type": tbl_width[0].get(qn("w:type")) if tbl_width else None,
                "border_element_count": tbl_borders,
                "header_row_count_in_sample": header_row_count,
                "sample_row_heights_mm": row_heights[:5],
                "sample_cell_widths_mm": column_widths[:8],
                "sample_cell_margins_twips": cell_margins[:3],
                "sample_style_count": len(sample_styles),
                "sample_styles": sample_styles[:3],
            }
        )
    return {"table_count": len(tables), "tables": tables}


def object_diagnostics(document: Any) -> dict[str, Any]:
    inline_shapes = getattr(document, "inline_shapes", [])
    inline_shape_details = []
    for idx in range(min(len(inline_shapes), 10)):
        shape = inline_shapes[idx]
        inline_shape_details.append(
            {
                "index": idx + 1,
                "type": enum_name(getattr(shape, "type", None)),
                "width_mm": mm_value(getattr(shape, "width", None)),
                "height_mm": mm_value(getattr(shape, "height", None)),
            }
        )
    body = document._element.body
    return {
        "inline_shape_count": len(inline_shapes),
        "inline_shape_samples": inline_shape_details,
        "drawing_element_count": safe_xpath_count(body, ".//w:drawing"),
        "legacy_pict_count": safe_xpath_count(body, ".//w:pict"),
        "text_box_count": safe_xpath_count(body, ".//w:txbxContent"),
        "watermark_like_shape_count": safe_xpath_count(body, ".//w:pict//v:shape"),
        "image_or_object_note": "Inline shapes may include images, seals, charts, or embedded objects; inspect manually before moving or resizing.",
    }


def style_system_diagnostics(document: Any) -> dict[str, Any]:
    type_names = {}
    builtin_count = 0
    custom_count = 0
    samples = []
    for style in document.styles:
        style_type = enum_name(style.type)
        type_names[style_type] = type_names.get(style_type, 0) + 1
        if style.builtin:
            builtin_count += 1
        else:
            custom_count += 1
        if len(samples) < 20:
            samples.append(
                {
                    "name": style.name,
                    "style_id": style.style_id,
                    "type": style_type,
                    "builtin": bool(style.builtin),
                }
            )
    return {
        "style_count": len(document.styles),
        "builtin_style_count": builtin_count,
        "custom_style_count": custom_count,
        "style_type_counts": type_names,
        "style_samples": samples,
    }


def special_state_diagnostics(document: Any) -> dict[str, Any]:
    body = document._element.body
    reltypes = [rel.reltype for rel in document.part.rels.values()]
    body_field_instructions = field_instructions(body)
    header_footer_field_instructions = [
        instruction
        for _, _, _, paragraph in iter_header_footer_paragraphs(document)
        for instruction in field_instructions(paragraph._element)
    ]
    all_field_instructions = [*body_field_instructions, *header_footer_field_instructions]
    field_diagnostics = {
        "body": {
            "instruction_count": len(body_field_instructions),
            "type_counts": field_type_counts(body_field_instructions),
            "category_counts": field_category_counts(body_field_instructions),
            "samples": field_instruction_details(body_field_instructions, scope="body"),
        },
        "headers_footers": {
            "instruction_count": len(header_footer_field_instructions),
            "type_counts": field_type_counts(header_footer_field_instructions),
            "category_counts": field_category_counts(header_footer_field_instructions),
            "samples": field_instruction_details(header_footer_field_instructions, scope="headers_footers"),
        },
        "all": {
            "instruction_count": len(all_field_instructions),
            "type_counts": field_type_counts(all_field_instructions),
            "category_counts": field_category_counts(all_field_instructions),
            "samples": [
                *field_instruction_details(body_field_instructions, scope="body", limit=10),
                *field_instruction_details(header_footer_field_instructions, scope="headers_footers", limit=10),
            ],
        },
        "notes": [
            "Field instructions are reported by type and category only; full instruction text is omitted by default.",
            "Field values may need updating in Word/WPS after formatting.",
        ],
    }
    update_risks = field_update_risks(field_diagnostics)
    return {
        "comments_relationship_count": len([rel for rel in reltypes if "comments" in rel]),
        "footnotes_relationship_count": len([rel for rel in reltypes if "footnotes" in rel]),
        "endnotes_relationship_count": len([rel for rel in reltypes if "endnotes" in rel]),
        "tracked_insertions": safe_xpath_count(body, ".//w:ins"),
        "tracked_deletions": safe_xpath_count(body, ".//w:del"),
        "field_simple_count": safe_xpath_count(body, ".//w:fldSimple"),
        "field_char_count": safe_xpath_count(body, ".//w:fldChar"),
        "body_field_instruction_count": len(body_field_instructions),
        "header_footer_field_instruction_count": len(header_footer_field_instructions),
        "all_field_instruction_count": len(all_field_instructions),
        "body_field_type_counts": field_type_counts(body_field_instructions),
        "header_footer_field_type_counts": field_type_counts(header_footer_field_instructions),
        "all_field_type_counts": field_type_counts(all_field_instructions),
        "field_diagnostics": field_diagnostics,
        "field_update_risks": update_risks,
        "field_update_risk_count": len(update_risks),
        "field_update_risk_category_counts": dict(Counter(risk["category"] for risk in update_risks)),
        "hyperlink_count": safe_xpath_count(body, ".//w:hyperlink"),
        "bookmark_count": safe_xpath_count(body, ".//w:bookmarkStart"),
        "footnote_reference_count": safe_xpath_count(body, ".//w:footnoteReference"),
        "endnote_reference_count": safe_xpath_count(body, ".//w:endnoteReference"),
        "hidden_text_run_count": safe_xpath_count(body, ".//w:rPr/w:vanish"),
    }


PARAGRAPH_CONTROL_COUNT_KEYS = [
    "paragraph_count",
    "keep_together_count",
    "keep_with_next_count",
    "page_break_before_count",
    "widow_control_count",
    "outline_level_count",
    "tab_stop_paragraph_count",
    "word_numbering_paragraph_count",
    "manual_numbering_paragraph_count",
    "manual_bullet_paragraph_count",
]


def paragraph_control_diagnostics(
    items: list[ParagraphItem],
    roles: list[dict[str, Any]],
) -> dict[str, Any]:
    counts = {key: 0 for key in PARAGRAPH_CONTROL_COUNT_KEYS}
    counts["paragraph_count"] = len(items)
    outline_levels: Counter[str] = Counter()
    manual_numbering: Counter[str] = Counter()
    manual_numbering_levels: Counter[str] = Counter()
    manual_numbering_sequence: list[dict[str, Any]] = []
    manual_hierarchy_warnings: list[dict[str, Any]] = []
    word_numbering_num_ids: Counter[str] = Counter()
    word_numbering_levels: Counter[str] = Counter()
    word_numbering_pairs: Counter[str] = Counter()
    word_numbering_samples: list[dict[str, Any]] = []
    samples = []
    previous_manual_level: int | None = None
    for item, role_info in zip(items, roles):
        paragraph = item.paragraph
        if paragraph is None:
            continue
        style = extract_paragraph_style(paragraph)
        controls = {
            key: style.get(key)
            for key in [
                "keep_together",
                "keep_with_next",
                "page_break_before",
                "widow_control",
                "outline_level",
                "numbering",
                "tab_stops",
            ]
            if style.get(key) is not None
        }
        if style.get("keep_together"):
            counts["keep_together_count"] += 1
        if style.get("keep_with_next"):
            counts["keep_with_next_count"] += 1
        if style.get("page_break_before"):
            counts["page_break_before_count"] += 1
        if style.get("widow_control"):
            counts["widow_control_count"] += 1
        if style.get("outline_level") is not None:
            counts["outline_level_count"] += 1
            outline_levels[str(style["outline_level"])] += 1
        if style.get("tab_stops"):
            counts["tab_stop_paragraph_count"] += 1
        numbering = style.get("numbering")
        if numbering:
            counts["word_numbering_paragraph_count"] += 1
            num_id = str(numbering.get("num_id") or "unknown")
            level = str(numbering.get("level") or "0")
            pair_key = f"numId={num_id};level={level}"
            word_numbering_num_ids[num_id] += 1
            word_numbering_levels[level] += 1
            word_numbering_pairs[pair_key] += 1
            if len(word_numbering_samples) < 12:
                word_numbering_samples.append(
                    {
                        "paragraph_index": role_info.get("index"),
                        "role": role_info.get("role"),
                        "hash": role_info.get("hash"),
                        "num_id": num_id,
                        "level": level,
                    }
                )

        manual_kind = manual_numbering_kind(item.text)
        if manual_kind:
            counts["manual_numbering_paragraph_count"] += 1
            manual_numbering[manual_kind] += 1
            level = manual_numbering_level(manual_kind)
            if level is not None:
                manual_numbering_levels[str(level)] += 1
                if previous_manual_level is not None and level - previous_manual_level > 1:
                    manual_hierarchy_warnings.append(
                        {
                            "paragraph_index": role_info.get("index"),
                            "hash": role_info.get("hash"),
                            "previous_level": previous_manual_level,
                            "current_level": level,
                            "warning": "Manual numbering level jumps by more than one.",
                        }
                    )
                previous_manual_level = level
            if len(manual_numbering_sequence) < 30:
                manual_numbering_sequence.append(
                    {
                        "paragraph_index": role_info.get("index"),
                        "role": role_info.get("role"),
                        "hash": role_info.get("hash"),
                        "kind": manual_kind,
                        "level": level,
                    }
                )
        normalized_text = normalize_text(item.text)
        bullet_like = bool(BULLET_RE.match(normalized_text)) and not SEPARATOR_RE.match(normalized_text)
        if bullet_like:
            counts["manual_bullet_paragraph_count"] += 1

        if len(samples) < 12 and (controls or manual_kind or bullet_like):
            sample = {
                "paragraph_index": role_info.get("index"),
                "role": role_info.get("role"),
                "hash": role_info.get("hash"),
                "controls": controls,
            }
            if manual_kind:
                sample["manual_numbering_kind"] = manual_kind
            if bullet_like:
                sample["manual_bullet_like"] = True
            samples.append(sample)

    return {
        "counts": counts,
        "outline_level_counts": dict(outline_levels),
        "manual_numbering_counts": dict(manual_numbering),
        "manual_numbering_level_counts": dict(manual_numbering_levels),
        "manual_numbering_sequence": manual_numbering_sequence,
        "manual_numbering_hierarchy_warnings": manual_hierarchy_warnings,
        "word_numbering_num_id_counts": dict(word_numbering_num_ids),
        "word_numbering_level_counts": dict(word_numbering_levels),
        "word_numbering_pair_counts": dict(word_numbering_pairs),
        "word_numbering_samples": word_numbering_samples,
        "samples": samples,
        "numbering_policy": (
            "Literal Chinese/Arabic markers are treated as existing text and styled as paragraphs. "
            "Word automatic numbering is detected and preserved; it is not regenerated automatically."
        ),
    }


def compare_paragraph_control_diagnostics(before: dict[str, Any], after: dict[str, Any]) -> dict[str, Any]:
    before_counts = before.get("counts", {})
    after_counts = after.get("counts", {})
    changes = []
    for key in PARAGRAPH_CONTROL_COUNT_KEYS:
        if before_counts.get(key) != after_counts.get(key):
            changes.append({"field": key, "before": before_counts.get(key), "after": after_counts.get(key)})
    return {
        "changed": bool(changes),
        "changes": changes,
        "before_counts": {key: before_counts.get(key) for key in PARAGRAPH_CONTROL_COUNT_KEYS},
        "after_counts": {key: after_counts.get(key) for key in PARAGRAPH_CONTROL_COUNT_KEYS},
    }


def role_style_summary(
    items: list[ParagraphItem],
    roles: list[dict[str, Any]],
    preset: str,
) -> dict[str, Any]:
    preset_styles = PRESETS[preset]
    summary: dict[str, Any] = {}
    for item, role_info in zip(items, roles):
        paragraph = item.paragraph
        if paragraph is None:
            continue
        role = role_info["role"]
        style = compact_style(extract_paragraph_style(paragraph))
        bucket = summary.setdefault(role, {"count": 0, "style_variants": Counter(), "samples": []})
        bucket["count"] += 1
        bucket["style_variants"][style_signature(style)] += 1
        if len(bucket["samples"]) < 3:
            expected = compact_style(preset_styles.get(role, preset_styles["body"]))
            diffs = []
            for key, exp in expected.items():
                actual = style.get(key)
                if actual is not None and exp is not None and actual != exp:
                    diffs.append({"field": key, "actual": actual, "expected": exp})
            bucket["samples"].append(
                {
                    "paragraph_index": role_info["index"],
                    "length": role_info["length"],
                    "hash": role_info["hash"],
                    "style": style,
                    "differences_from_preset": diffs,
                }
            )

    normalized: dict[str, Any] = {}
    for role, data in summary.items():
        variants = [
            {"count": count, "style": json.loads(signature)}
            for signature, count in data["style_variants"].most_common()
        ]
        normalized[role] = {
            "count": data["count"],
            "style_variant_count": len(variants),
            "style_variants": variants[:5],
            "samples": data["samples"],
        }
    return normalized


def build_format_diagnostics(
    document: Any,
    items: list[ParagraphItem],
    roles: list[dict[str, Any]],
    preset: str,
) -> dict[str, Any]:
    role_summary = role_style_summary(items, roles, preset)
    warnings = []
    inconsistent_roles = [
        role for role, data in role_summary.items() if data.get("style_variant_count", 0) > 1
    ]
    if inconsistent_roles:
        warnings.append("Style variants detected within roles: " + ", ".join(sorted(inconsistent_roles)))
    if document.tables:
        warnings.append("Tables detected; review table borders, widths, and header rows manually if strict formatting is required.")
    if getattr(document, "inline_shapes", []):
        warnings.append("Images or inline objects detected; preserve seals/images unless the user explicitly requests repositioning.")
    special_state = special_state_diagnostics(document)
    if special_state.get("tracked_insertions") or special_state.get("tracked_deletions"):
        warnings.append("Tracked changes detected; preserve revisions unless the user explicitly requests accepting or rejecting them.")
    if special_state.get("field_simple_count") or special_state.get("field_char_count"):
        warnings.append("Field codes detected; verify generated tables of contents, dates, or references after formatting.")
    for risk in special_state.get("field_update_risks", []):
        warnings.append(risk["warning"])

    return {
        "page": page_diagnostics(document, preset),
        "headers_footers": header_footer_diagnostics(document),
        "role_style_summary": role_summary,
        "paragraph_controls": paragraph_control_diagnostics(items, roles),
        "tables": table_diagnostics(document),
        "objects": object_diagnostics(document),
        "style_system": style_system_diagnostics(document),
        "special_state": special_state,
        "diagnostic_warnings": warnings,
        "diagnostic_scope": [
            "page setup",
            "headers and footers",
            "paragraph roles",
            "font family and size",
            "bold italic underline strikethrough color highlight",
            "indentation spacing alignment pagination controls tab stops numbering",
            "heading hierarchy",
            "tables",
            "inline images or objects",
            "style system",
            "comments tracked changes fields hyperlinks footnotes endnotes",
        ],
    }


def build_coverage(
    document: Any,
    roles: list[dict[str, Any]],
    counts: dict[str, int],
    mode: str,
    diagnostics: dict[str, Any] | None,
    format_actions: dict[str, Any] | None = None,
) -> dict[str, Any]:
    format_actions = format_actions or {}
    detected_roles = sorted(role for role in counts if role != "empty")
    formatted: list[dict[str, Any]] = []
    preserved: list[dict[str, Any]] = []
    diagnosed_only: list[dict[str, Any]] = []
    needs_review: list[dict[str, Any]] = []

    if mode == "diagnose-only":
        diagnosed_only.extend(
            [
                {"area": "page_setup", "status": "diagnosed_only"},
                {"area": "headers_footers", "status": "diagnosed_only"},
                {"area": "paragraph_roles", "status": "diagnosed_only", "roles": detected_roles},
                {"area": "typography_and_paragraph_layout", "status": "diagnosed_only"},
                {"area": "paragraph_controls", "status": "diagnosed_only"},
                {"area": "tables", "status": "diagnosed_only", "count": len(document.tables)},
                {
                    "area": "inline_images_or_objects",
                    "status": "diagnosed_only",
                    "count": len(getattr(document, "inline_shapes", [])),
                },
            ]
        )
    else:
        formatted.extend(
            [
                {"area": "page_setup", "status": "formatted"},
                {"area": "paragraph_roles", "status": "formatted", "roles": detected_roles},
                {"area": "typography_and_paragraph_layout", "status": "formatted"},
                {
                    "area": "paragraph_controls",
                    "status": "formatted",
                    "note": "Safe pagination controls, widow control, and outline levels are applied to detected existing paragraphs.",
                },
            ]
        )
        table_actions = format_actions.get("tables", [])
        if document.tables and table_actions:
            formatted.append(
                {
                    "area": "table_structure",
                    "status": "formatted",
                    "count": len(document.tables),
                    "note": "Existing table structure was normalized because --format-tables was used.",
                }
            )
        if document.tables:
            formatted.append(
                {
                    "area": "table_text",
                    "status": "formatted",
                    "count": len(document.tables),
                    "note": "Table paragraph text received body-style formatting where directly accessible.",
                }
            )
        if document.tables and not table_actions:
            preserved.append(
                {
                    "area": "table_structure",
                    "status": "preserved",
                    "count": len(document.tables),
                    "note": "Borders, widths, row heights, merged cells, and table layout are preserved unless explicitly handled.",
                }
            )

    inline_shape_count = len(getattr(document, "inline_shapes", []))
    if inline_shape_count:
        preserved.append(
            {
                "area": "inline_images_or_objects",
                "status": "preserved",
                "count": inline_shape_count,
                "note": "Images, seals, charts, or embedded inline objects are preserved unless the user explicitly asks for repositioning.",
                }
            )

    paragraph_controls = (
        diagnostics.get("paragraph_controls")
        if diagnostics
        else paragraph_control_diagnostics(document_items(document), roles)
    )
    paragraph_control_counts = paragraph_controls.get("counts", {})
    if paragraph_control_counts.get("word_numbering_paragraph_count"):
        target_bucket = diagnosed_only if mode == "diagnose-only" else preserved
        target_bucket.append(
            {
                "area": "word_automatic_numbering",
                "status": "diagnosed_only" if mode == "diagnose-only" else "preserved",
                "count": paragraph_control_counts.get("word_numbering_paragraph_count"),
                "note": "Existing Word numbering is detected and kept; automatic numbering definitions are not regenerated.",
            }
        )
    if paragraph_control_counts.get("manual_numbering_paragraph_count"):
        target_bucket = diagnosed_only if mode == "diagnose-only" else formatted
        target_bucket.append(
            {
                "area": "manual_numbering_text",
                "status": "diagnosed_only" if mode == "diagnose-only" else "formatted",
                "count": paragraph_control_counts.get("manual_numbering_paragraph_count"),
                "manual_numbering_counts": paragraph_controls.get("manual_numbering_counts", {}),
                "note": "Existing numbering markers such as 一、（一） and 1. are treated as text and styled with their paragraph role.",
            }
        )
    if paragraph_control_counts.get("manual_bullet_paragraph_count"):
        target_bucket = diagnosed_only if mode == "diagnose-only" else preserved
        target_bucket.append(
            {
                "area": "manual_bullets",
                "status": "diagnosed_only" if mode == "diagnose-only" else "preserved",
                "count": paragraph_control_counts.get("manual_bullet_paragraph_count"),
                "note": "Manual bullet markers are preserved as existing text.",
            }
        )
    if paragraph_control_counts.get("tab_stop_paragraph_count"):
        target_bucket = diagnosed_only if mode == "diagnose-only" else preserved
        target_bucket.append(
            {
                "area": "tab_stops",
                "status": "diagnosed_only" if mode == "diagnose-only" else "preserved",
                "count": paragraph_control_counts.get("tab_stop_paragraph_count"),
                "note": "Existing tab stops are preserved unless a supported formatter rule changes them.",
            }
        )

    header_footer = diagnostics.get("headers_footers") if diagnostics else header_footer_diagnostics(document)
    page_field_count = sum(
        int(section.get("page_field_count", 0))
        for section in header_footer.get("sections", [])
    )
    if page_field_count:
        target_bucket = diagnosed_only if mode == "diagnose-only" else formatted
        target_bucket.append(
            {
                "area": "page_numbers",
                "status": "diagnosed_only" if mode == "diagnose-only" else "formatted",
                "page_field_count": page_field_count,
            }
        )
    header_footer_with_text = [
        section
        for section in header_footer.get("sections", [])
        if section.get("has_header_text") or section.get("has_footer_text")
    ]
    if header_footer_with_text and mode != "diagnose-only":
        preserved.append(
            {
                "area": "headers_footers",
                "status": "preserved",
                "section_count": len(header_footer_with_text),
                "note": "Existing header/footer text is preserved by default.",
            }
        )

    not_detected = [
        {"area": "paragraph_role", "status": "not_detected", "role": role}
        for role in COMMON_REPORT_ROLES
        if role not in counts
    ]
    if not page_field_count:
        not_detected.append({"area": "page_numbers", "status": "not_detected"})

    for role_info in roles:
        if role_info.get("role") == "needs_review" or role_info.get("warnings"):
            needs_review.append(
                {
                    "area": "paragraph_role",
                    "status": "needs_review",
                    "paragraph_index": role_info.get("index"),
                    "role": role_info.get("role"),
                    "confidence": role_info.get("confidence"),
                    "reason_codes": role_info.get("reason_codes", []),
                    "warnings": role_info.get("warnings", []),
                }
            )

    unsupported = list(LIMITED_FORMAT_AREAS)
    coverage = {
        "formatted": formatted,
        "preserved": preserved,
        "diagnosed_only": diagnosed_only,
        "not_detected": not_detected,
        "unsupported": unsupported,
        "needs_review": needs_review,
    }
    coverage["summary"] = {
        key: len(value)
        for key, value in coverage.items()
        if isinstance(value, list)
    }
    return coverage


def classify_items(items: list[ParagraphItem], preset: str) -> tuple[list[dict[str, Any]], dict[str, int]]:
    total = len(items)
    trailing_date_indexes = {
        idx
        for idx, item in enumerate(items)
        if idx >= total - 3 and DATE_RE.fullmatch(normalize_text(item.text))
    }
    candidate_rows: list[list[RoleCandidate]] = []
    seen_separator = False
    for idx, item in enumerate(items):
        context = {
            "seen_separator_before": seen_separator,
            "has_trailing_date_after": any(date_idx > idx for date_idx in trailing_date_indexes),
        }
        candidates = paragraph_role_candidates(item, idx, total, preset, context)
        candidate_rows.append(candidates)
        if any(candidate.role == "separator" for candidate in candidates):
            seen_separator = True

    selected_main_title_index = None
    for idx, candidates in enumerate(candidate_rows):
        if any(candidate.role == "main_title" and candidate.confidence >= 0.8 for candidate in candidates):
            selected_main_title_index = idx
            break

    selected_article_title_index = None

    roles: list[dict[str, Any]] = []
    counts: dict[str, int] = {}
    for idx, (item, candidates) in enumerate(zip(items, candidate_rows)):
        decision = decide_role(candidates, selected_main_title_index, selected_article_title_index, idx)
        role = decision.role
        counts[role] = counts.get(role, 0) + 1
        roles.append(
            {
                "index": idx + 1,
                "role": role,
                "confidence": round(decision.confidence, 2),
                "reason_codes": decision.reason_codes,
                "length": len(item.text.strip()),
                "hash": text_hash(item.text),
                "warnings": decision.warnings,
            }
        )
    return roles, counts


def format_document(
    document: Any,
    items: list[ParagraphItem],
    preset: str,
) -> tuple[list[dict[str, Any]], dict[str, int], list[str]]:
    preset_styles = PRESETS[preset]
    page_style = dict(preset_styles["_page"])
    apply_page_setup(document, page_style)

    roles, counts = classify_items(items, preset)
    for item, info in zip(items, roles):
        paragraph = item.paragraph
        if paragraph is None:
            continue
        role = info["role"]
        base_style = preset_styles.get(role, preset_styles["body"])
        apply_style(paragraph, base_style)

    for paragraph in iter_table_paragraphs(document):
        if paragraph.text.strip():
            apply_style(paragraph, preset_styles["body"])

    notes = []
    return roles, counts, notes


def default_output_path(input_path: Path | None, input_name: str) -> Path:
    if input_path is not None:
        return input_path.with_name(input_path.stem + "_formatted.docx")
    return Path(input_name).with_suffix("").with_name(Path(input_name).stem + "_formatted.docx")


def candidate_output_path(output_path: Path) -> Path:
    for attempt in range(100):
        suffix = f".candidate-{os.getpid()}-{attempt}.docx"
        candidate = output_path.with_name(f".{output_path.name}{suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"Could not allocate a candidate output path for {output_path}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Format existing .docx files with the fixed 行文检查版式 without changing text.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  format_document.py input.docx --output output.docx --report report.json\n"
            "  format_document.py input.docx --diagnose-only --report diagnostics.json\n\n"
            "Notes:\n"
            "  - Only existing .docx input is supported.\n"
            "  - Missing page numbers are added only with --add-page-numbers.\n"
            "  - Table structure is normalized only with --format-tables.\n"
        ),
    )
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("--output", help="Output .docx path")
    parser.add_argument(
        "--preset",
        choices=["checklist"],
        default="checklist",
        help="Formatting preset to apply; only the fixed 行文检查版式 is supported",
    )
    parser.add_argument("--report", help="Optional JSON report path")
    parser.add_argument("--diagnose-only", action="store_true", help="Diagnose structure and formatting; do not save .docx")
    parser.add_argument("--add-page-numbers", action="store_true", help="Explicitly add page-number fields to footers when missing")
    parser.add_argument("--format-tables", action="store_true", help="Explicitly normalize existing table structure formatting")
    parser.add_argument("--include-text-in-report", action="store_true", help="Include full paragraph text in JSON report")
    return parser.parse_args()


def print_cli_status(report: dict[str, Any], report_path: Path | None) -> None:
    summary = report.get("summary", {})
    user_message = summary.get("user_message", {})
    title = user_message.get("title")
    body = user_message.get("body")
    if title:
        print(title, file=sys.stderr)
    if body:
        print(body, file=sys.stderr)
    for label, key in [
        ("能做什么", "can_do"),
        ("当前状态", "current_state"),
        ("下一步", "next_action"),
    ]:
        value = user_message.get(key)
        if value:
            print(f"{label}: {value}", file=sys.stderr)
    output = summary.get("output")
    if output:
        print(f"输出文件: {output}", file=sys.stderr)
    if report_path:
        print(f"报告文件: {report_path}", file=sys.stderr)


def main() -> int:
    args = parse_args()

    require_docx()

    input_path = Path(args.input).resolve()
    source_name = str(input_path)
    suffix = input_path.suffix.lower()

    if not input_path.exists():
        raise SystemExit(f"Input file not found: {input_path}. Please provide an existing .docx file.")
    if suffix == ".docx":
        try:
            document = Document(str(input_path))
        except Exception as exc:
            raise SystemExit(f"Could not open .docx file: {input_path}. Make sure it is a valid, unencrypted Word document. ({exc})")
        items = document_items(document)
    else:
        input_type = suffix or "<no extension>"
        raise SystemExit(f"Unsupported input type: {input_type}. Use .docx only; this skill does not convert Markdown, text, PDF, or images.")
    before_text_fingerprint = document_text_fingerprint(document)
    before_page_diagnostics = page_diagnostics(document, args.preset)
    before_roles, _ = classify_items(items, args.preset)
    before_paragraph_controls = paragraph_control_diagnostics(items, before_roles)

    diagnose_only = args.diagnose_only
    requested_output_path = None if diagnose_only else Path(args.output).resolve() if args.output else default_output_path(input_path, input_path.name).resolve()
    actual_output_path = None
    output_withheld = False
    exit_code = 0
    report_document = document

    if diagnose_only:
        roles, counts = classify_items(items, args.preset)
        style_notes = []
        diagnostics = build_format_diagnostics(document, items, roles, args.preset)
        mode = "diagnose-only"
        page_number_actions = []
        table_actions = []
        after_text_fingerprint = document_text_fingerprint(document)
        after_page_diagnostics = page_diagnostics(document, args.preset)
        after_paragraph_controls = paragraph_control_diagnostics(items, roles)
    else:
        roles, counts, style_notes = format_document(document, items, args.preset)
        page_number_actions = format_existing_page_numbers(document, args.preset)
        if args.add_page_numbers:
            page_number_actions.extend(add_page_numbers(document, args.preset))
        table_actions = format_table_structures(document, args.preset) if args.format_tables else []
        diagnostics = None
        mode = "format-only"
        assert requested_output_path is not None
        candidate_path = candidate_output_path(requested_output_path)
        document.save(str(candidate_path))
        report_document = Document(str(candidate_path))
        candidate_items = document_items(report_document)
        after_text_fingerprint = document_text_fingerprint(report_document)
        after_page_diagnostics = page_diagnostics(report_document, args.preset)
        after_paragraph_controls = paragraph_control_diagnostics(candidate_items, roles)

    generated_page_number_actions = [
        action for action in page_number_actions if action.get("action") == "added_page_number"
    ]
    content_preservation = compare_text_fingerprints(
        before_text_fingerprint,
        after_text_fingerprint,
        generated_layout_elements=generated_page_number_actions,
    )
    format_changes = {
        "page": compare_page_diagnostics(before_page_diagnostics, after_page_diagnostics),
        "paragraph_controls": compare_paragraph_control_diagnostics(
            before_paragraph_controls,
            after_paragraph_controls,
        ),
    }
    format_actions = {"page_numbers": page_number_actions, "tables": table_actions}
    coverage = build_coverage(report_document, roles, counts, mode, diagnostics, format_actions=format_actions)

    if not diagnose_only:
        assert requested_output_path is not None
        if content_status(content_preservation) == "changed":
            output_withheld = True
            exit_code = 1
            try:
                candidate_path.unlink()
            except FileNotFoundError:
                pass
        else:
            candidate_path.replace(requested_output_path)
            actual_output_path = requested_output_path

    report = build_report(
        source_name=source_name,
        preset=args.preset,
        mode=mode,
        roles=roles,
        counts=counts,
        output=actual_output_path,
        requested_output=requested_output_path,
        output_withheld=output_withheld,
        style_notes=style_notes,
        diagnostics=diagnostics,
        include_text=args.include_text_in_report,
        items=items,
        content_preservation=content_preservation,
        coverage=coverage,
        format_changes=format_changes,
        format_actions=format_actions,
    )
    report_path = Path(args.report).resolve() if args.report else None
    if args.report or diagnose_only or output_withheld:
        write_report(report, report_path)
    if IS_CLI_ENTRYPOINT and (not args.report or output_withheld):
        print_cli_status(report, report_path)
    return exit_code


if __name__ == "__main__":
    raise SystemExit(main())
