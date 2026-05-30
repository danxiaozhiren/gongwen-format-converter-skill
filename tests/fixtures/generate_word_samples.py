#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate minimal .docx fixtures for Word-format coverage checks."""

from __future__ import annotations

import base64
from contextlib import contextmanager
import shutil
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory, gettempdir
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Mm, Pt

try:
    import fcntl
except ImportError:
    fcntl = None


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "word_samples"
IMAGE_DIR = OUTPUT_DIR / "_assets"
LOCK_PATH = Path(gettempdir()) / "gongwen-word-fixtures.lock"

PNG_1X1_RED = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMB/ax"
    "G8xkAAAAASUVORK5CYII="
)


@contextmanager
def fixture_generation_lock():
    LOCK_PATH.parent.mkdir(parents=True, exist_ok=True)
    with LOCK_PATH.open("w", encoding="utf-8") as lock_file:
        if fcntl is not None:
            fcntl.flock(lock_file.fileno(), fcntl.LOCK_EX)
        try:
            yield
        finally:
            if fcntl is not None:
                fcntl.flock(lock_file.fileno(), fcntl.LOCK_UN)


def reset_output_dir() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for path in OUTPUT_DIR.glob("*.docx"):
        path.unlink()
    IMAGE_DIR.mkdir(exist_ok=True)
    (IMAGE_DIR / "seal-placeholder.png").write_bytes(PNG_1X1_RED)


def set_run_font(run, font: str = "仿宋_GB2312", size_pt: int = 16, bold: bool | None = None) -> None:
    run.font.name = font
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)


def add_paragraph(document: Document, text: str, *, align=None, font="仿宋_GB2312", size=16, bold=None):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, font=font, size_pt=size, bold=bold)
    return paragraph


def basic_page_setup(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(32)
        section.bottom_margin = Mm(28)
        section.left_margin = Mm(30)
        section.right_margin = Mm(30)


def add_page_field(paragraph, instruction: str = "PAGE") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run._r.append(instr)

    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    paragraph.add_run("1")

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def add_toc_field(paragraph) -> None:
    add_page_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')


def add_bookmark(paragraph, name: str, bookmark_id: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def set_direct_numbering(paragraph, num_id: str = "5", level: str = "0") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), level)
    num_pr.append(ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.append(num_id_el)


def add_textbox(paragraph, text: str) -> None:
    pict = parse_xml(
        f"""
        <w:r {nsdecls("w")} xmlns:v="urn:schemas-microsoft-com:vml">
          <w:pict>
            <v:shape id="TextBox1" type="#_x0000_t202" style="width:180pt;height:45pt">
              <v:textbox inset="6pt,6pt,6pt,6pt">
                <w:txbxContent>
                  <w:p>
                    <w:r>
                      <w:t>{text}</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </v:textbox>
            </v:shape>
          </w:pict>
        </w:r>
        """
    )
    paragraph._p.append(pict)


def sample_basic_formal() -> Path:
    document = Document()
    basic_page_setup(document)
    add_paragraph(document, "关于开展格式检查工作的通知", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=22)
    add_paragraph(document, "各部门：")
    add_paragraph(document, "为规范内部材料格式，现就有关事项通知如下。")
    add_paragraph(document, "一、总体要求", font="黑体")
    add_paragraph(document, "各部门应当保持正文内容稳定，只对版式和样式进行调整。")
    add_paragraph(document, "（一）加强审核", font="楷体_GB2312")
    add_paragraph(document, "1. 明确责任。材料提交前应完成格式自查。")
    add_paragraph(document, "附件：格式检查清单")
    add_paragraph(document, "办公室", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(document, "2026年5月30日", align=WD_ALIGN_PARAGRAPH.RIGHT)
    path = OUTPUT_DIR / "01-basic-formal.docx"
    document.save(path)
    return path


def sample_table() -> Path:
    document = Document()
    basic_page_setup(document)
    add_paragraph(document, "表格材料样例", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=22)
    table = document.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["事项", "责任部门", "完成情况"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    table.cell(1, 0).text = "材料汇总"
    table.cell(1, 1).text = "办公室"
    table.cell(1, 2).text = "进行中"
    table.cell(2, 0).text = "格式检查"
    table.cell(2, 1).text = "综合处"
    table.cell(2, 2).text = "已完成"
    table.cell(3, 0).merge(table.cell(3, 1))
    table.cell(3, 0).text = "备注"
    table.cell(3, 2).text = "含合并单元格"
    path = OUTPUT_DIR / "02-table.docx"
    document.save(path)
    return path


def sample_header_footer_page() -> Path:
    document = Document()
    basic_page_setup(document)
    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.header.paragraphs[0].text = "内部资料"
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    add_page_field(footer, "PAGE")
    footer.add_run(" 页 / 共 ")
    add_page_field(footer, "NUMPAGES")
    footer.add_run(" 页")
    add_paragraph(document, "页眉页脚和页码样例", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=22)
    add_paragraph(document, "本样例用于检查页眉、页脚、PAGE 字段和首页不同设置。")
    document.add_section(WD_SECTION_START.NEW_PAGE)
    add_paragraph(document, "第二节内容", font="黑体")
    add_paragraph(document, "第二节用于检查多节文档的页面处理。")
    path = OUTPUT_DIR / "03-header-footer-page.docx"
    document.save(path)
    return path


def sample_image_seal() -> Path:
    document = Document()
    basic_page_setup(document)
    add_paragraph(document, "图片和印章占位样例", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=22)
    add_paragraph(document, "下方红色方块用于模拟印章类图片对象，格式化器默认应保留其位置和大小。")
    document.add_picture(str(IMAGE_DIR / "seal-placeholder.png"), width=Mm(22), height=Mm(22))
    add_paragraph(document, "图片下方正文继续保持原有顺序。")
    path = OUTPUT_DIR / "04-image-seal.docx"
    document.save(path)
    return path


def sample_textbox_shape() -> Path:
    document = Document()
    basic_page_setup(document)
    add_paragraph(document, "文本框和形状样例", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=22)
    add_paragraph(document, "下方包含一个简单文本框 OOXML，用于检查 text_box_count 诊断。")
    paragraph = document.add_paragraph()
    add_textbox(paragraph, "文本框内文字")
    add_paragraph(document, "文本框之后的正文。")
    path = OUTPUT_DIR / "05-textbox-shape.docx"
    document.save(path)
    return path


def sample_auto_numbering() -> Path:
    document = Document()
    basic_page_setup(document)
    add_paragraph(document, "自动编号样例", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=22)
    for text in ["自动编号第一项", "自动编号第二项", "自动编号第三项"]:
        paragraph = document.add_paragraph(text, style="List Number")
        set_direct_numbering(paragraph)
        for run in paragraph.runs:
            set_run_font(run)
    add_paragraph(document, "一、手写编号也应作为正文文本保留", font="黑体")
    path = OUTPUT_DIR / "06-auto-numbering.docx"
    document.save(path)
    return path


def sample_toc_fields() -> Path:
    document = Document()
    basic_page_setup(document)
    add_paragraph(document, "目录和域代码样例", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=22)
    toc = document.add_paragraph()
    add_toc_field(toc)
    date_paragraph = document.add_paragraph("生成日期：")
    add_page_field(date_paragraph, 'DATE \\@ "yyyy年M月d日"')
    ref_paragraph = document.add_paragraph("交叉引用标题：")
    add_page_field(ref_paragraph, "REF SectionOne \\h")
    page_ref_paragraph = document.add_paragraph("交叉引用页码：")
    add_page_field(page_ref_paragraph, "PAGEREF SectionOne \\h")
    heading = add_paragraph(document, "一、第一部分", font="黑体")
    add_bookmark(heading, "SectionOne", "1")
    add_paragraph(document, "第一部分正文。")
    add_paragraph(document, "二、第二部分", font="黑体")
    add_paragraph(document, "第二部分正文。")
    path = OUTPUT_DIR / "07-toc-fields.docx"
    document.save(path)
    return path


def sample_comments_revisions_base() -> Path:
    document = Document()
    basic_page_setup(document)
    add_paragraph(document, "批注和修订样例", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=22)
    add_paragraph(document, "本段将被插入批注锚点。")
    add_paragraph(document, "本段后续会通过 OOXML 注入修订标记。")
    path = OUTPUT_DIR / "08-comments-revisions.docx"
    document.save(path)
    return path


def next_rel_id(rels_root: ET.Element) -> str:
    max_id = 0
    for rel in rels_root:
        rid = rel.attrib.get("Id", "")
        if rid.startswith("rId") and rid[3:].isdigit():
            max_id = max(max_id, int(rid[3:]))
    return f"rId{max_id + 1}"


def add_comment_footnote_endnote_and_revision(path: Path) -> None:
    with TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        with zipfile.ZipFile(path, "r") as archive:
            archive.extractall(temp_path)

        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
            "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
        }
        document_xml = temp_path / "word" / "document.xml"
        tree = ET.parse(document_xml)
        root = tree.getroot()
        body = root.find("w:body", ns)
        assert body is not None
        paragraphs = body.findall("w:p", ns)
        comment_target = paragraphs[1]
        revision_target = paragraphs[2]

        comment_start = ET.Element(qn("w:commentRangeStart"), {qn("w:id"): "0"})
        comment_end = ET.Element(qn("w:commentRangeEnd"), {qn("w:id"): "0"})
        comment_run = ET.Element(qn("w:r"))
        comment_ref = ET.SubElement(comment_run, qn("w:commentReference"), {qn("w:id"): "0"})
        comment_target.insert(0, comment_start)
        comment_target.append(comment_end)
        comment_target.append(comment_run)

        footnote_run = ET.Element(qn("w:r"))
        ET.SubElement(footnote_run, qn("w:footnoteReference"), {qn("w:id"): "1"})
        endnote_run = ET.Element(qn("w:r"))
        ET.SubElement(endnote_run, qn("w:endnoteReference"), {qn("w:id"): "1"})
        revision_target.append(footnote_run)
        revision_target.append(endnote_run)

        inserted = ET.Element(
            qn("w:ins"),
            {qn("w:id"): "1", qn("w:author"): "fixture", qn("w:date"): "2026-05-30T00:00:00Z"},
        )
        ins_run = ET.SubElement(inserted, qn("w:r"))
        ET.SubElement(ins_run, qn("w:t")).text = "插入的修订文字"
        deleted = ET.Element(
            qn("w:del"),
            {qn("w:id"): "2", qn("w:author"): "fixture", qn("w:date"): "2026-05-30T00:00:00Z"},
        )
        del_run = ET.SubElement(deleted, qn("w:r"))
        ET.SubElement(del_run, qn("w:delText")).text = "删除的修订文字"
        revision_target.append(inserted)
        revision_target.append(deleted)
        tree.write(document_xml, encoding="UTF-8", xml_declaration=True)

        comments_xml = temp_path / "word" / "comments.xml"
        comments_xml.write_text(
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="0" w:author="fixture" w:date="2026-05-30T00:00:00Z">
    <w:p><w:r><w:t>批注占位，不用于内容输出。</w:t></w:r></w:p>
  </w:comment>
</w:comments>
""",
            encoding="utf-8",
        )
        (temp_path / "word" / "footnotes.xml").write_text(
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:id="-1" w:type="separator"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>
  <w:footnote w:id="0" w:type="continuationSeparator"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>
  <w:footnote w:id="1"><w:p><w:r><w:t>脚注占位文本。</w:t></w:r></w:p></w:footnote>
</w:footnotes>
""",
            encoding="utf-8",
        )
        (temp_path / "word" / "endnotes.xml").write_text(
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:endnote w:id="-1" w:type="separator"><w:p><w:r><w:separator/></w:r></w:p></w:endnote>
  <w:endnote w:id="0" w:type="continuationSeparator"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:endnote>
  <w:endnote w:id="1"><w:p><w:r><w:t>尾注占位文本。</w:t></w:r></w:p></w:endnote>
</w:endnotes>
""",
            encoding="utf-8",
        )

        content_types_xml = temp_path / "[Content_Types].xml"
        content_tree = ET.parse(content_types_xml)
        content_root = content_tree.getroot()
        for part_name, content_type in [
            ("/word/comments.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"),
            ("/word/footnotes.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"),
            ("/word/endnotes.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"),
        ]:
            override_exists = any(
                element.attrib.get("PartName") == part_name
                for element in content_root.findall("ct:Override", ns)
            )
            if not override_exists:
                ET.SubElement(
                    content_root,
                    "{http://schemas.openxmlformats.org/package/2006/content-types}Override",
                    {
                        "PartName": part_name,
                        "ContentType": content_type,
                    },
                )
        content_tree.write(content_types_xml, encoding="UTF-8", xml_declaration=True)

        rels_xml = temp_path / "word" / "_rels" / "document.xml.rels"
        rels_tree = ET.parse(rels_xml)
        rels_root = rels_tree.getroot()
        for rel_type, target in [
            ("http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments", "comments.xml"),
            ("http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes", "footnotes.xml"),
            ("http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes", "endnotes.xml"),
        ]:
            rel_exists = any(rel.attrib.get("Type") == rel_type for rel in rels_root)
            if not rel_exists:
                ET.SubElement(
                    rels_root,
                    "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship",
                    {
                        "Id": next_rel_id(rels_root),
                        "Type": rel_type,
                        "Target": target,
                    },
                )
        rels_tree.write(rels_xml, encoding="UTF-8", xml_declaration=True)

        rebuilt = path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(rebuilt, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for file_path in temp_path.rglob("*"):
                if file_path.is_file():
                    archive.write(file_path, file_path.relative_to(temp_path).as_posix())
        shutil.move(rebuilt, path)


def generate_all() -> list[Path]:
    with fixture_generation_lock():
        reset_output_dir()
        paths = [
            sample_basic_formal(),
            sample_table(),
            sample_header_footer_page(),
            sample_image_seal(),
            sample_textbox_shape(),
            sample_auto_numbering(),
            sample_toc_fields(),
            sample_comments_revisions_base(),
        ]
        add_comment_footnote_endnote_and_revision(paths[-1])
        return paths


def main() -> int:
    paths = generate_all()
    for path in paths:
        print(path.relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
