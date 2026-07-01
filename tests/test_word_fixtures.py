#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Unittest entrypoint for Word fixture smoke checks."""

from __future__ import annotations

import importlib.util
import json
from pathlib import Path
import subprocess
import sys
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn

from tests.fixtures import check_word_samples


ROOT = Path(__file__).resolve().parents[1]
FORMATTER = ROOT / "skills" / "format-xingwen-word" / "scripts" / "format_document.py"
EVALS = ROOT / "skills" / "format-xingwen-word" / "evals" / "evals.json"


def write_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(str(path))


def run_font_attrs(run):
    rpr = run._element.rPr
    rfonts = rpr.rFonts if rpr is not None else None
    if rfonts is None:
        return {}
    return {
        "eastAsia": rfonts.get(qn("w:eastAsia")),
        "ascii": rfonts.get(qn("w:ascii")),
        "hAnsi": rfonts.get(qn("w:hAnsi")),
    }


def load_formatter_module():
    module_name = "format_document_under_test"
    formatter_dir = str(FORMATTER.parent)
    if formatter_dir not in sys.path:
        sys.path.insert(0, formatter_dir)
    spec = importlib.util.spec_from_file_location(module_name, FORMATTER)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load formatter module from {FORMATTER}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


class WordFixtureSmokeTest(unittest.TestCase):
    def test_generated_samples_emit_expected_diagnostics(self) -> None:
        self.assertEqual(check_word_samples.main(), 0)

    def test_cli_default_uses_checklist_layout(self) -> None:
        with TemporaryDirectory(prefix="xingwen-checklist-layout-") as temp_dir:
            temp_path = Path(temp_dir)
            source = temp_path / "checklist-layout.docx"
            output = temp_path / "checklist-layout-output.docx"
            report = temp_path / "checklist-layout.json"
            write_docx(
                source,
                [
                    "关于开展默认版式检查工作的通知",
                    "各部门：",
                    "正文内容包含 ABC123，用于检查数字和字母字体。",
                    "一、一级标题",
                    "（一）二级标题",
                    "1. 三级标题",
                    "（1）四级标题",
                ],
            )

            subprocess.run(
                [
                    sys.executable,
                    str(FORMATTER),
                    str(source),
                    "--output",
                    str(output),
                    "--add-page-numbers",
                    "--report",
                    str(report),
                ],
                cwd=ROOT,
                check=True,
            )

            data = json.loads(report.read_text(encoding="utf-8"))
            self.assertEqual(data["preset"], "checklist")
            self.assertEqual(data["summary"]["format_source"], "xingwen-checklist 行文出手前对照检查事项")
            self.assertEqual(data["summary"]["output"], str(output.resolve()))
            self.assertEqual(data["summary"]["requested_output"], str(output.resolve()))
            self.assertFalse(data["summary"]["output_withheld"])
            self.assertIn("已完成", data["summary"]["user_message"]["title"])
            self.assertIn("已有 .docx", data["summary"]["user_message"]["can_do"])
            self.assertIn("输出文件已生成", data["summary"]["user_message"]["current_state"])
            self.assertIn("最终视觉确认", data["summary"]["user_message"]["next_action"])
            self.assertEqual(data["summary"]["user_message"]["output"], str(output.resolve()))
            self.assertEqual(data["summary"]["content_status"], "changed_by_explicit_layout_addition")
            self.assertIn("page_setup", data["summary"]["formatted_areas"])
            self.assertEqual(data["summary"]["document_grid_reference"]["line_count"], 22)
            self.assertEqual(data["summary"]["document_grid_reference"]["char_count"], 28)
            self.assertTrue(data["summary"]["recommended_actions"])
            self.assertEqual(data["role_counts"]["heading_4"], 1)

            document = Document(str(output))
            section = document.sections[0]
            self.assertAlmostEqual(section.top_margin.mm, 37, places=1)
            self.assertAlmostEqual(section.bottom_margin.mm, 35, places=1)
            self.assertAlmostEqual(section.left_margin.mm, 27, places=1)
            self.assertAlmostEqual(section.right_margin.mm, 27, places=1)

            paragraphs = {paragraph.text: paragraph for paragraph in document.paragraphs if paragraph.text}
            title = paragraphs["关于开展默认版式检查工作的通知"]
            self.assertAlmostEqual(title.paragraph_format.line_spacing.pt, 36, places=1)
            self.assertAlmostEqual(title.paragraph_format.space_after.pt, 30, places=1)
            title_run = title.runs[0]
            self.assertEqual(title_run.font.size.pt, 22)
            self.assertFalse(title_run.font.bold)
            self.assertEqual(run_font_attrs(title_run)["eastAsia"], "方正小标宋简体")
            self.assertEqual(run_font_attrs(title_run)["ascii"], "Times New Roman")

            body_run = paragraphs["正文内容包含 ABC123，用于检查数字和字母字体。"].runs[0]
            self.assertAlmostEqual(
                paragraphs["正文内容包含 ABC123，用于检查数字和字母字体。"].paragraph_format.line_spacing.pt,
                30,
                places=1,
            )
            self.assertEqual(run_font_attrs(body_run)["eastAsia"], "仿宋_GB2312")
            self.assertEqual(run_font_attrs(body_run)["ascii"], "Times New Roman")

            heading_2_run = paragraphs["（一）二级标题"].runs[0]
            self.assertTrue(heading_2_run.font.bold)
            self.assertEqual(run_font_attrs(heading_2_run)["eastAsia"], "楷体_GB2312")

            heading_4 = paragraphs["（1）四级标题"]
            self.assertAlmostEqual(heading_4.paragraph_format.line_spacing.pt, 30, places=1)
            self.assertEqual(run_font_attrs(heading_4.runs[0])["eastAsia"], "仿宋_GB2312")

            footer_run = section.footer.paragraphs[0].runs[0]
            self.assertEqual(footer_run.font.size.pt, 14)
            self.assertEqual(run_font_attrs(footer_run)["eastAsia"], "宋体")

    def test_content_preservation_failure_withholds_output(self) -> None:
        with TemporaryDirectory(prefix="gongwen-content-failure-") as temp_dir:
            temp_path = Path(temp_dir)
            source = temp_path / "source.docx"
            output = temp_path / "source-output.docx"
            report = temp_path / "source.json"
            write_docx(
                source,
                [
                    "关于开展内容保持检查的通知",
                    "各部门：",
                    "正文内容不得被修改。",
                ],
            )

            formatter = load_formatter_module()
            original_format_document = formatter.format_document

            def mutating_format_document(document, items, preset):
                roles, counts, notes = original_format_document(document, items, preset)
                document.paragraphs[0].add_run("错误新增文本")
                return roles, counts, notes

            argv = [
                "format_document.py",
                str(source),
                "--output",
                str(output),
                "--report",
                str(report),
            ]
            with patch.object(formatter, "format_document", side_effect=mutating_format_document), patch.object(sys, "argv", argv):
                exit_code = formatter.main()

            self.assertEqual(exit_code, 1)
            self.assertFalse(output.exists())
            self.assertFalse(list(temp_path.glob(".*candidate-*.docx")))
            data = json.loads(report.read_text(encoding="utf-8"))
            self.assertIsNone(data["output"])
            self.assertEqual(data["requested_output"], str(output.resolve()))
            self.assertTrue(data["output_withheld"])
            self.assertEqual(data["summary"]["status"], "failed_content_preservation")
            self.assertTrue(data["summary"]["output_withheld"])
            self.assertIn("已停止交付", data["summary"]["user_message"]["title"])
            self.assertIn("正式输出未交付", data["summary"]["user_message"]["current_state"])
            self.assertIn("重新运行", data["summary"]["user_message"]["next_action"])
            self.assertIsNone(data["summary"]["user_message"]["output"])
            self.assertTrue(data["content_preservation"]["text_changed"])

    def test_format_tables_preserves_table_text(self) -> None:
        with TemporaryDirectory(prefix="gongwen-format-tables-") as temp_dir:
            temp_path = Path(temp_dir)
            source = temp_path / "table-source.docx"
            output = temp_path / "table-output.docx"
            report = temp_path / "table-output.json"
            document = Document()
            document.add_paragraph("表格材料样例")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "项目"
            table.cell(0, 1).text = "数值"
            table.cell(1, 0).text = "收入"
            table.cell(1, 1).text = "100"
            document.save(str(source))

            subprocess.run(
                [
                    sys.executable,
                    str(FORMATTER),
                    str(source),
                    "--output",
                    str(output),
                    "--format-tables",
                    "--report",
                    str(report),
                ],
                cwd=ROOT,
                check=True,
            )

            data = json.loads(report.read_text(encoding="utf-8"))
            self.assertTrue(output.exists())
            self.assertFalse(data["output_withheld"])
            self.assertFalse(data["content_preservation"]["text_changed"])
            self.assertIn("table_structure", data["summary"]["formatted_areas"])
            result = Document(str(output))
            self.assertEqual(result.tables[0].cell(0, 0).text, "项目")
            self.assertEqual(result.tables[0].cell(1, 1).text, "100")

    def test_diagnose_only_does_not_create_output(self) -> None:
        with TemporaryDirectory(prefix="gongwen-diagnose-only-") as temp_dir:
            temp_path = Path(temp_dir)
            source = temp_path / "diagnose.docx"
            output = temp_path / "diagnose-output.docx"
            report = temp_path / "diagnose.json"
            write_docx(source, ["关于诊断格式的通知", "各部门：", "正文。"])

            subprocess.run(
                [
                    sys.executable,
                    str(FORMATTER),
                    str(source),
                    "--diagnose-only",
                    "--output",
                    str(output),
                    "--report",
                    str(report),
                ],
                cwd=ROOT,
                check=True,
            )

            data = json.loads(report.read_text(encoding="utf-8"))
            self.assertFalse(output.exists())
            self.assertIsNone(data["output"])
            self.assertIsNone(data["requested_output"])
            self.assertFalse(data["output_withheld"])
            self.assertEqual(data["summary"]["mode"], "diagnose-only")

    def test_cli_rejects_non_docx_input(self) -> None:
        with TemporaryDirectory(prefix="xingwen-word-only-") as temp_dir:
            temp_path = Path(temp_dir)
            source = temp_path / "source.txt"
            output = temp_path / "source-output.docx"
            report = temp_path / "source.json"
            source.write_text("这不是 Word 文件。", encoding="utf-8")

            completed = subprocess.run(
                [
                    sys.executable,
                    str(FORMATTER),
                    str(source),
                    "--output",
                    str(output),
                    "--report",
                    str(report),
                ],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            )

            self.assertNotEqual(completed.returncode, 0)
            self.assertIn("Use .docx only", completed.stderr)
            self.assertIn("does not convert Markdown, text, PDF, or images", completed.stderr)
            self.assertFalse(output.exists())
            self.assertFalse(report.exists())

    def test_cli_reports_friendly_missing_and_corrupt_docx_errors(self) -> None:
        with TemporaryDirectory(prefix="xingwen-friendly-errors-") as temp_dir:
            temp_path = Path(temp_dir)
            missing = temp_path / "missing.docx"
            missing_completed = subprocess.run(
                [
                    sys.executable,
                    str(FORMATTER),
                    str(missing),
                ],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertNotEqual(missing_completed.returncode, 0)
            self.assertIn("Input file not found", missing_completed.stderr)
            self.assertIn("existing .docx file", missing_completed.stderr)

            corrupt = temp_path / "corrupt.docx"
            corrupt.write_bytes(b"not a valid docx package")
            corrupt_completed = subprocess.run(
                [
                    sys.executable,
                    str(FORMATTER),
                    str(corrupt),
                ],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertNotEqual(corrupt_completed.returncode, 0)
            self.assertIn("Could not open .docx file", corrupt_completed.stderr)
            self.assertIn("valid, unencrypted Word document", corrupt_completed.stderr)

    def test_cli_without_report_prints_short_user_status(self) -> None:
        with TemporaryDirectory(prefix="xingwen-cli-status-") as temp_dir:
            temp_path = Path(temp_dir)
            source = temp_path / "status.docx"
            output = temp_path / "status_formatted.docx"
            write_docx(source, ["关于检查交互提示的通知", "各部门：", "正文内容。"])

            completed = subprocess.run(
                [
                    sys.executable,
                    str(FORMATTER),
                    str(source),
                ],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            )

            self.assertEqual(completed.returncode, 0)
            self.assertTrue(output.exists())
            self.assertIn("已完成", completed.stderr)
            self.assertIn("能做什么", completed.stderr)
            self.assertIn("当前状态", completed.stderr)
            self.assertIn("下一步", completed.stderr)
            self.assertIn("输出文件", completed.stderr)
            self.assertIn(str(output.resolve()), completed.stderr)

    def test_role_detection_does_not_treat_doc_number_as_main_title(self) -> None:
        with TemporaryDirectory(prefix="gongwen-role-detection-") as temp_dir:
            temp_path = Path(temp_dir)
            source = temp_path / "doc-number-first.docx"
            report = temp_path / "doc-number-first.json"
            write_docx(source, ["中办发〔2024〕1号", "关于开展角色识别检查的通知", "各部门：", "正文内容。"])

            subprocess.run(
                [
                    sys.executable,
                    str(FORMATTER),
                    str(source),
                    "--diagnose-only",
                    "--report",
                    str(report),
                ],
                cwd=ROOT,
                check=True,
            )

            data = json.loads(report.read_text(encoding="utf-8"))
            self.assertEqual(data["paragraphs"][0]["role"], "metadata")
            self.assertEqual(data["paragraphs"][1]["role"], "main_title")
            self.assertIn("metadata_pattern", data["paragraphs"][0]["reason_codes"])
            self.assertIn("early_title_like_non_metadata", data["paragraphs"][1]["reason_codes"])

    def test_evals_have_machine_readable_expected_signals(self) -> None:
        data = json.loads(EVALS.read_text(encoding="utf-8"))
        for eval_case in data["evals"]:
            run = eval_case.get("run")
            self.assertIsInstance(run, dict, msg=f"eval {eval_case.get('id')} missing runnable config")
            self.assertIn(
                run.get("kind"),
                {"formatter", "formatter_mutation", "mock_agent_policy"},
                msg=f"eval {eval_case.get('id')} has unsupported run.kind",
            )
            signals = eval_case.get("expected_signals")
            self.assertIsInstance(signals, list, msg=f"eval {eval_case.get('id')} missing expected_signals")
            self.assertTrue(signals, msg=f"eval {eval_case.get('id')} has no expected_signals")
            for signal in signals:
                self.assertIsInstance(signal.get("path"), str)
                assertions = {"equals", "minimum", "contains"} & set(signal)
                self.assertEqual(len(assertions), 1, msg=f"eval {eval_case.get('id')} signal must have one assertion")

if __name__ == "__main__":
    unittest.main()
